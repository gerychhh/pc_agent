from __future__ import annotations

import re
import time
from dataclasses import dataclass
from typing import Any
from urllib.parse import quote_plus

import pyautogui
import pygetwindow


YT_HOME = "https://www.youtube.com/"
YT_SEARCH_URL = "https://www.youtube.com/results?search_query={query}"
DEFAULT_VIDEO_URL = "https://www.youtube.com/watch?v=dQw4w9WgXcQ"

DEFAULT_ABOUT_PARAGRAPHS = [
    "Я — локальный ассистент управления Windows‑ПК через PowerShell и Python.",
    "Я умею открывать сайты и программы, создавать документы (DOCX/XLSX/PPTX), а также управлять YouTube и Блокнотом.",
    "Если команда не сработала — я покажу ошибку и попробую исправить её в умном режиме.",
]


@dataclass
class Action:
    language: str
    script: str
    updates: dict[str, Any] | None = None
    name: str | None = None


def _focus_window(keywords: tuple[str, ...]) -> bool:
    for window in pygetwindow.getAllWindows():
        title = (window.title or "").lower()
        if any(keyword in title for keyword in keywords):
            if window.isMinimized:
                window.restore()
            window.activate()
            time.sleep(0.2)
            return True
    return False


def _youtube_window_open() -> bool:
    return _focus_window(("youtube", "chrome", "firefox", "edge"))


def _notepad_window_open() -> bool:
    return _focus_window(("notepad", "блокнот"))


def _powershell_start_process(target: str) -> str:
    return f'Start-Process "{target}"'


def generate_word_doc_ps(text: str, filename: str = "me.docx") -> str:
    """Word COM вариант (требует установленный Microsoft Word).

    В основном проекте предпочтительнее python-docx (оно надёжнее и не зависит от Word),
    но этот генератор оставлен как альтернативный вариант.
    """
    safe_text = text.replace('"', '""')
    return (
        "$desktop = [Environment]::GetFolderPath('Desktop')\n"
        f"$path = Join-Path $desktop '{filename}'\n"
        "$word = New-Object -ComObject Word.Application\n"
        "$word.Visible = $false\n"
        "$doc = $word.Documents.Add()\n"
        f"$doc.Content.Text = @\"\n{safe_text}\n\"@\n"
        "$doc.SaveAs([ref]$path, [ref]16)\n"
        "$doc.Close()\n"
        "$word.Quit()\n"
        'Write-Host "SAVED: $path"\n'
    )


def _extract_docx_filename(user_text: str) -> str | None:
    """Пытаемся вытащить имя файла, если пользователь явно указал.

    Примеры:
    - "... назови файл report.docx"
    - "... сохрани как о_себе.docx"
    """
    match = re.search(
        r"(?:назови|имя|как|сохрани\s+как)\s+([\w\-а-яА-Я_]+\.docx)",
        user_text,
        re.IGNORECASE,
    )
    if match:
        name = match.group(1).strip()
        if name.lower().endswith(".docx"):
            return name
    return None


def _extract_docx_text(user_text: str) -> str | None:
    """Пытаемся вытащить текст для документа из фразы пользователя."""
    match = re.search(
        r"(?:впиши\s+туда|напиши\s+туда|помести\s+туда|и\s+напиши)\s*[:\-]?\s*(.+)$",
        user_text,
        re.IGNORECASE,
    )
    if match:
        text = match.group(1).strip()
        if text:
            return text
    return None


def _python_create_docx(title: str, paragraphs: list[str], filename: str) -> str:
    """Генерирует безопасный Python-скрипт для создания реального .docx через python-docx."""
    safe_title = repr(title)
    safe_paragraphs = repr(paragraphs)
    safe_filename = repr(filename)

    return (
        "import os\n"
        "from pathlib import Path\n"
        "from docx import Document\n"
        "\n"
        "def pick_desktop() -> Path:\n"
        "    candidates = []\n"
        "    for key in ('OneDriveConsumer', 'OneDrive'):\n"
        "        val = os.environ.get(key)\n"
        "        if val:\n"
        "            candidates.append(Path(val) / 'Desktop')\n"
        "    candidates.append(Path.home() / 'Desktop')\n"
        "    user = os.environ.get('USERPROFILE')\n"
        "    if user:\n"
        "        candidates.append(Path(user) / 'Desktop')\n"
        "    for p in candidates:\n"
        "        if p.exists():\n"
        "            return p\n"
        "    return Path.home()\n"
        "\n"
        f"title = {safe_title}\n"
        f"paragraphs = {safe_paragraphs}\n"
        f"filename = {safe_filename}\n"
        "desktop = pick_desktop()\n"
        "path = desktop / filename\n"
        "doc = Document()\n"
        "if title:\n"
        "    doc.add_heading(title, level=1)\n"
        "for p in paragraphs:\n"
        "    if p:\n"
        "        doc.add_paragraph(str(p))\n"
        "doc.save(str(path))\n"
        "print(f'SAVED: {path}')\n"
    )


def match_skill(user_text: str, state: dict[str, Any]) -> Action | str | None:
    lowered = user_text.lower().strip()

    # =============================
    # DOCX / Word document creation (железный скилл)
    # =============================
    docx_keywords = ("docx", "ворд", "word", "word файл", "ворд файл", "документ ворд")
    create_verbs = ("создай", "создать", "сделай", "сделать", "сгенерируй")

    is_docx_request = any(k in lowered for k in docx_keywords) and (
        any(v in lowered for v in create_verbs) or "файл" in lowered or "документ" in lowered
    )

    # Частый кейс: пользователь говорит просто "word файл" после ошибки
    if lowered in ("word файл", "ворд файл", "word", "ворд"):
        is_docx_request = True

    if is_docx_request and "txt" not in lowered:
        filename = _extract_docx_filename(user_text) or "о_себе.docx"
        explicit_text = _extract_docx_text(user_text)
        paragraphs = [explicit_text] if explicit_text else DEFAULT_ABOUT_PARAGRAPHS
        script = _python_create_docx("О себе", paragraphs, filename)
        return Action(language="python", script=script, name="create_docx")

    # =============================
    # YouTube
    # =============================
    if "ютуб" in lowered or "youtube" in lowered or "на ютуб" in lowered:
        if lowered.startswith("открой ютуб") or lowered == "открой ютуб":
            return Action(
                language="powershell",
                script=_powershell_start_process(YT_HOME),
                updates={"active_url": YT_HOME},
                name="youtube_open",
            )

        search_match = re.search(r"найди на (?:ютубе|ютуб|youtube)\s+(.+)", user_text, re.IGNORECASE)
        if search_match:
            query = search_match.group(1).strip()
            url = YT_SEARCH_URL.format(query=quote_plus(query))
            return Action(
                language="powershell",
                script=_powershell_start_process(url),
                updates={"active_url": url},
                name="youtube_search",
            )

        if "включи видео" in lowered or "открой видео" in lowered:
            active_url = state.get("active_url") or ""
            url = DEFAULT_VIDEO_URL
            if "results?search_query=" in active_url:
                url = DEFAULT_VIDEO_URL
            return Action(
                language="powershell",
                script=_powershell_start_process(url),
                updates={"active_url": url},
                name="youtube_video",
            )

    if any(term in lowered for term in ("пауза", "стоп", "продолжи")) and (
        "ютуб" in lowered or "youtube" in lowered or state.get("active_url")
    ):
        if not _youtube_window_open():
            return "Открой YouTube сначала."
        script = _python_youtube_control("play_pause", 1)
        return Action(language="python", script=script, name="youtube_play_pause")

    if (
        ("перемотай" in lowered and "вперед" in lowered)
        or ("впер" in lowered and ("сек" in lowered or "мин" in lowered))
    ) and ("ютуб" in lowered or state.get("active_url")):
        if not _youtube_window_open():
            return "Открой YouTube сначала."
        seconds = _extract_seconds(lowered) or 10
        steps = max(1, round(seconds / 10))
        return Action(language="python", script=_python_youtube_control("seek_forward_10", steps), name="youtube_seek")

    if "назад" in lowered and ("ютуб" in lowered or state.get("active_url")):
        if not _youtube_window_open():
            return "Открой YouTube сначала."
        seconds = _extract_seconds(lowered) or 10
        steps = max(1, round(seconds / 10))
        return Action(language="python", script=_python_youtube_control("seek_back_10", steps), name="youtube_back")

    if "выключи звук" in lowered and ("ютуб" in lowered or state.get("active_url")):
        if not _youtube_window_open():
            return "Открой YouTube сначала."
        return Action(language="python", script=_python_youtube_control("mute_toggle", 1), name="youtube_mute")

    if "громче" in lowered and ("ютуб" in lowered or state.get("active_url")):
        if not _youtube_window_open():
            return "Открой YouTube сначала."
        return Action(language="python", script=_python_youtube_control("volume_up", 3), name="youtube_volume_up")

    if "тише" in lowered and ("ютуб" in lowered or state.get("active_url")):
        if not _youtube_window_open():
            return "Открой YouTube сначала."
        return Action(language="python", script=_python_youtube_control("volume_down", 3), name="youtube_volume_down")

    if "на весь экран" in lowered and ("ютуб" in lowered or state.get("active_url")):
        if not _youtube_window_open():
            return "Открой YouTube сначала."
        return Action(language="python", script=_python_youtube_control("fullscreen_toggle", 1), name="youtube_fullscreen")

    # =============================
    # Notepad / Word open/close/write
    # =============================
    if "открой блокнот" in lowered:
        return Action(
            language="powershell",
            script="Start-Process notepad.exe",
            updates={"active_app": "notepad"},
            name="open_notepad",
        )

    if "открой ворд" in lowered or "открой word" in lowered:
        return Action(
            language="powershell",
            script="Start-Process winword.exe",
            updates={"active_app": "winword.exe"},
            name="open_word",
        )

    if "закрой блокнот" in lowered:
        return Action(
            language="powershell",
            script="Stop-Process -Name notepad -ErrorAction SilentlyContinue",
            updates={"active_app": "notepad"},
            name="close_notepad",
        )

    write_match = re.search(r"напиши в (?:этом )?блокноте\s+(.+)", user_text, re.IGNORECASE)
    if write_match:
        if not _notepad_window_open():
            return "Открой блокнот сначала."
        text = write_match.group(1).strip()
        return Action(language="python", script=_python_type_in_notepad(text), name="type_notepad")

    # =============================
    # Open URL by domain (quick)
    # =============================
    url_match = re.search(r"открой\s+([a-z0-9.-]+\.[a-z]{2,})(/\S*)?$", lowered)
    if url_match:
        url = url_match.group(1)
        if not url.startswith(("http://", "https://")):
            url = f"https://{url}"
        return Action(language="powershell", script=_powershell_start_process(url), updates={"active_url": url})

    # =============================
    # Small app aliases
    # =============================
    app_aliases = {"калькулятор": "calc.exe", "calc": "calc.exe"}
    for alias, exe in app_aliases.items():
        if f"открой {alias}" in lowered:
            return Action(
                language="powershell",
                script=f"Start-Process {exe}",
                updates={"active_app": exe},
                name=f"open_{alias}",
            )

    return None


def _extract_seconds(text: str) -> int | None:
    match = re.search(r"(\d+)\s*(сек|секунд|секунды|с)", text)
    if match:
        return int(match.group(1))
    match = re.search(r"(\d+)\s*(мин|минута|минуты|минут)", text)
    if match:
        return int(match.group(1)) * 60
    if "мин" in text and not re.search(r"\d", text):
        return 60
    return None


def _python_youtube_control(action: str, steps: int) -> str:
    return (
        "import time\n"
        "import pyautogui\n"
        "import pygetwindow\n"
        "\n"
        "def focus():\n"
        "    for window in pygetwindow.getAllWindows():\n"
        "        title = (window.title or \"\").lower()\n"
        "        if any(key in title for key in (\"youtube\", \"chrome\", \"firefox\", \"edge\")):\n"
        "            if window.isMinimized:\n"
        "                window.restore()\n"
        "            window.activate()\n"
        "            time.sleep(0.2)\n"
        "            return True\n"
        "    return False\n"
        "\n"
        "if not focus():\n"
        "    raise RuntimeError(\"Не найдено окно браузера/YouTube. Открой YouTube сначала.\")\n"
        "pyautogui.FAILSAFE = True\n"
        f"action = \"{action}\"\n"
        f"steps = {steps}\n"
        "def press_many(key, count):\n"
        "    for _ in range(max(1, count)):\n"
        "        pyautogui.press(key)\n"
        "        time.sleep(0.05)\n"
        "if action == \"play_pause\":\n"
        "    pyautogui.press(\"k\")\n"
        "elif action == \"seek_forward_10\":\n"
        "    press_many(\"l\", steps)\n"
        "elif action == \"seek_back_10\":\n"
        "    press_many(\"j\", steps)\n"
        "elif action == \"mute_toggle\":\n"
        "    pyautogui.press(\"m\")\n"
        "elif action == \"volume_up\":\n"
        "    press_many(\"up\", steps)\n"
        "elif action == \"volume_down\":\n"
        "    press_many(\"down\", steps)\n"
        "elif action == \"fullscreen_toggle\":\n"
        "    pyautogui.press(\"f\")\n"
    )


def _python_type_in_notepad(text: str) -> str:
    safe = text.replace("\\", "\\\\").replace("\"", "\\\"")
    return (
        "import time\n"
        "import pyautogui\n"
        "import pygetwindow\n"
        "\n"
        "def focus():\n"
        "    for window in pygetwindow.getAllWindows():\n"
        "        title = (window.title or \"\").lower()\n"
        "        if \"notepad\" in title or \"блокнот\" in title:\n"
        "            if window.isMinimized:\n"
        "                window.restore()\n"
        "            window.activate()\n"
        "            time.sleep(0.2)\n"
        "            return True\n"
        "    return False\n"
        "\n"
        "if not focus():\n"
        "    raise RuntimeError(\"Не найдено окно блокнота. Открой блокнот сначала.\")\n"
        "pyautogui.typewrite(\"" + safe + "\", interval=0.01)\n"
    )
