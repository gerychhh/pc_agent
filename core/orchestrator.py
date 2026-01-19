from __future__ import annotations

import hashlib
import re
from typing import Any

from .config import DEBUG, MAX_RETRIES, TIMEOUT_SEC
from .executor import run_pip_install, run_powershell, run_python
from .llm_client import LLMClient
from .logger import SessionLogger
from .policy import RiskLevel, assess_risk, confirm_if_needed
from .state import load_state
from .youtube import focus_browser_window, open_first_video_from_search, open_search, open_youtube_home


SYSTEM_PROMPT = """
Ты — локальный ассистент управления Windows-ПК через PowerShell и Python.

Отвечай КРАТКО.

ГЛАВНОЕ ПРАВИЛО:
Если пользователь просит СДЕЛАТЬ действие на компьютере (открыть/запустить/перейти/создать/скачать/найти/настроить/закрыть) —
ТЫ ОБЯЗАН вернуть КОД ДЛЯ ВЫПОЛНЕНИЯ.
Не спорь, не объясняй, не “я не это”, а просто делай.

ФОРМАТ ОТВЕТА:
1) Если действие НЕ нужно (вопрос/объяснение) — 1–3 предложения, БЕЗ кода.
2) Если действие НУЖНО — верни РОВНО ОДИН блок кода и НИЧЕГО больше:
   - либо ```powershell ... ```
   - либо ```python ... ```
Вне code block не пиши ни одного символа.

ПРАВИЛА ДЛЯ ДЕЙСТВИЙ:
- По умолчанию используй PowerShell.
- Скрипт должен быть ПОЛНЫМ и ИСПОЛНЯЕМЫМ.
- Если нужно открыть сайт — используй:
  Start-Process "https://...."
- Если пользователь пишет "ютуб / youtube" — открывай https://www.youtube.com/

ОКРУЖЕНИЕ:
- ОС: Windows, shell: PowerShell.
- Python доступен как `python`.
- Для Desktop:
  - PowerShell: `$desktop = [Environment]::GetFolderPath('Desktop')`
  - Python: `from pathlib import Path; desktop = Path.home() / "Desktop"`
  
ФАЙЛЫ И ФОРМАТЫ (ОЧЕНЬ ВАЖНО):
Если пользователь просит создать файл определённого типа (docx, xlsx, pdf, png, mp3 и т.д.) —
файл должен быть РЕАЛЬНОГО ФОРМАТА, а не просто текстом с таким расширением.

Запрещено делать так:
- open("file.docx","w").write("...")  # это НЕ docx
- open("file.pdf","w").write("...")   # это НЕ pdf

Делай правильно:
- .docx → только через библиотеку python-docx:
  from docx import Document; doc = Document(); doc.add_paragraph(...); doc.save("file.docx")
- .xlsx → только через openpyxl или pandas ExcelWriter
- .pdf  → только через reportlab (или другой генератор PDF)
- .json/.txt/.md/.csv → можно через обычный open(...,"w") (это текстовые форматы)
- картинки (.png/.jpg) → генерируй/сохраняй как изображение, не текст

ВЫБОР ЯЗЫКА ДЛЯ ФАЙЛОВ:
- .docx / .xlsx / .pdf → всегда Python (правильные библиотеки)
- открыть сайт / запустить программу / команды ОС → PowerShell

Если библиотека недоступна или формат неясен — задай 1 уточняющий вопрос или предложи правильный способ.

ВАЖНО ПРО ЯЗЫК СКРИПТА:
Если в решении используется Python-код или Python-библиотеки (например python-docx, openpyxl, reportlab и т.д.) —
ТОГДА скрипт обязан быть в формате ```python``` и выполняться как Python.
Никогда не вставляй Python-код внутрь PowerShell.

PowerShell используй только для команд Windows (Start-Process, cd, dir, New-Item и т.д.).

АКТИВНЫЙ ФАЙЛ (ACTIVE_FILE):
- Если ACTIVE_FILE указан и это .docx, то для изменения шрифта используй python-docx.
- Открывай существующий документ: Document(active_file)
- Применяй шрифт стилю Normal и всем runs
- Сохраняй в тот же файл
- Запрещено создавать новый Document() при наличии ACTIVE_FILE

ПРАВИЛА ДЛЯ PYTHON:
- Запрещены однострочники с ';'
- Скрипт Python должен начинаться с нужных импортов (первая непустая строка — import/from)
- .docx → только python-docx, .xlsx → только openpyxl, .pdf → только reportlab
- Никогда не создавай docx/xlsx/pdf через open(...,"w")

БЕЗОПАСНОСТЬ:
- Не делай разрушительные действия.
- Если запрос опасный или неясный — задай 1 короткий уточняющий вопрос.
""".strip()




def sanitize_assistant_text(text: str) -> str:
    if not text:
        return ""
    cleaned = text.replace("[TOOL_RESULT]", "").replace("[END_TOOL_RESULT]", "")
    filtered_lines: list[str] = []
    for line in cleaned.splitlines():
        if any(marker in line for marker in ("<|channel|>", "<|constrain|>", "<|message|>")):
            continue
        filtered_lines.append(line)
    return "\n".join(filtered_lines).strip()


class Orchestrator:
    def __init__(self, history_limit: int = 4) -> None:
        self.client = LLMClient()
        self.logger = SessionLogger()
        self.history_limit = history_limit
        self.system_message = {"role": "system", "content": SYSTEM_PROMPT}
        self.reset()

    def reset(self) -> None:
        self.history: list[dict[str, Any]] = []

    @staticmethod
    def _route_intent(user_input: str) -> tuple[str, str | None]:
        text = user_input.lower()
        file_keywords = (
            "создай файл",
            "документ",
            "ворд",
            "docx",
            "таблицу",
            "xlsx",
            "pdf",
        )
        os_keywords = (
            "открой",
            "запусти",
            "перейди",
            "включи",
            "закрой",
        )
        if any(keyword in text for keyword in file_keywords):
            return "file", "python"
        if any(keyword in text for keyword in os_keywords):
            return "os_action", "powershell"
        return "unknown", None

    def _build_messages(
        self,
        user_input: str,
        stateless: bool,
        system_override: str | None = None,
    ) -> list[dict[str, Any]]:
        messages: list[dict[str, Any]] = [self.system_message]
        if system_override:
            messages.append({"role": "system", "content": system_override})
        if not stateless and self.history:
            messages.extend(self.history[-self.history_limit :])
        messages.append({"role": "user", "content": user_input})
        return messages

    @staticmethod
    def _format_action(action: dict[str, str]) -> str:
        return f"```{action['language']}\n{action['script']}\n```"

    @staticmethod
    def _parse_duration_seconds(text: str) -> int | None:
        lowered = text.lower()
        if "мину" in lowered and not re.search(r"\d", lowered):
            return 60
        match = re.search(r"(\d+)\s*(сек|секунд|секунды|с|мин|минут|минута|минуты)", lowered)
        if match:
            value = int(match.group(1))
            unit = match.group(2)
            if unit.startswith("мин"):
                return value * 60
            return value
        digits = re.search(r"(\d+)", lowered)
        if digits:
            return int(digits.group(1))
        return None

    def _handle_youtube_command(self, user_input: str, active_url: str | None) -> str | None:
        lowered = user_input.lower()
        has_youtube_keyword = any(term in lowered for term in ("ютуб", "ютубе", "youtube", "на ютуб"))
        has_active_youtube = bool(active_url and ("youtube" in active_url or "youtu.be" in active_url))
        if not has_youtube_keyword and not has_active_youtube:
            return None

        if "закрой" in lowered and has_youtube_keyword:
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"close_tab\")"
            return f"```python\n{script}\n```"

        if "открой" in lowered and ("ютуб" in lowered or "youtube" in lowered):
            action = open_youtube_home()
            return self._format_action(action)

        search_match = re.search(
            r"найди на (?:ютубе|ютуб|youtube)\s+(.+)", user_input, re.IGNORECASE
        )
        if search_match:
            action = open_search(search_match.group(1).strip())
            return self._format_action(action)

        play_match = re.search(r"(?:включи|открой) видео\s+(.+)", user_input, re.IGNORECASE)
        if play_match:
            action = open_first_video_from_search(play_match.group(1).strip())
            return self._format_action(action)

        if any(word in lowered for word in ("пауза", "останови")):
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"play_pause\")"
            return f"```python\n{script}\n```"

        if any(word in lowered for word in ("продолжи", "воспроизведение")):
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"play_pause\")"
            return f"```python\n{script}\n```"

        if "перемотай" in lowered or "назад" in lowered:
            seconds = self._parse_duration_seconds(lowered) or 10
            steps = max(1, round(seconds / 10))
            if "назад" in lowered:
                action = "seek_back_10"
            else:
                action = "seek_forward_10"
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = f"from core.youtube import control\ncontrol(\"{action}\", {steps})"
            return f"```python\n{script}\n```"

        if "громче" in lowered:
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"volume_up\", 3)"
            return f"```python\n{script}\n```"

        if "тише" in lowered:
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"volume_down\", 3)"
            return f"```python\n{script}\n```"

        if "выключи звук" in lowered or "mute" in lowered or "включи звук" in lowered:
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"mute_toggle\")"
            return f"```python\n{script}\n```"

        if "на весь экран" in lowered or "полный экран" in lowered:
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"fullscreen_toggle\")"
            return f"```python\n{script}\n```"

        if "следующее видео" in lowered or "следующий" in lowered:
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"next_video\")"
            return f"```python\n{script}\n```"

        if "предыдущее видео" in lowered or "предыдущ" in lowered:
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"prev_video\")"
            return f"```python\n{script}\n```"

        if "субтитр" in lowered:
            if not focus_browser_window():
                return "Открой YouTube сначала: 'открой ютуб'."
            script = "from core.youtube import control\ncontrol(\"captions_toggle\")"
            return f"```python\n{script}\n```"

        return None

    @staticmethod
    def _is_file_operation(user_input: str) -> bool:
        triggers = (
            "измени",
            "поменяй",
            "сделай",
            "шрифт",
            "формат",
            "добавь",
            "вставь",
            "удали",
            "перепиши",
            "переведи",
            "сохрани",
            "экспорт",
        )
        lowered = user_input.lower()
        return any(trigger in lowered for trigger in triggers)

    @staticmethod
    def _has_explicit_file(user_input: str) -> bool:
        lowered = user_input.lower()
        return any(ext in lowered for ext in (".docx", ".txt", ".xlsx", ".pdf", ".png", ".jpg"))

    @staticmethod
    def _format_active_context(active_file: str, active_type: str | None) -> str:
        type_value = active_type or ""
        return f"[CONTEXT]\nACTIVE_FILE: {active_file}\nACTIVE_TYPE: {type_value}\n[/CONTEXT]\n"

    @staticmethod
    def _extract_script(content: str) -> tuple[str | None, str | None]:
        if not content:
            return None, None
        python_match = re.search(r"```python\s*(.*?)```", content, re.DOTALL | re.IGNORECASE)
        if python_match:
            return "python", python_match.group(1).strip()
        ps_match = re.search(r"```powershell\s*(.*?)```", content, re.DOTALL | re.IGNORECASE)
        if ps_match:
            return "powershell", ps_match.group(1).strip()
        return None, None

    @staticmethod
    def _is_command_text(content: str) -> bool:
        if not content:
            return False
        trimmed = content.strip()
        command_prefixes = (
            "get-",
            "set-",
            "start-",
            "new-",
            "remove-",
            "copy-",
            "move-",
            "invoke-",
            "add-",
            "python ",
            "pip ",
            "dir",
            "cd ",
            "echo ",
            "notepad",
            "start ",
        )
        lowered = trimmed.lower()
        return lowered.startswith(command_prefixes)

    def _run_script(self, language: str, script: str) -> dict[str, Any]:
        if language == "python":
            return run_python(script, TIMEOUT_SEC)
        return run_powershell(script, TIMEOUT_SEC)

    def _log_debug(self, label: str, value: str) -> None:
        if DEBUG:
            print(f"[{label}] {value}")

    @staticmethod
    def _script_hash(script: str) -> str:
        return hashlib.sha256(script.encode("utf-8")).hexdigest()[:8]

    @staticmethod
    def _extract_missing_module(stderr: str) -> str | None:
        match = re.search(r"No module named ['\"]([^'\"]+)['\"]", stderr)
        if match:
            return match.group(1)
        return None

    @staticmethod
    def _map_package(module_name: str) -> str | None:
        mapping = {
            "docx": "python-docx",
            "win32com": "pywin32",
        }
        return mapping.get(module_name)

    @staticmethod
    def _powershell_script_incomplete(script: str) -> bool:
        stripped = script.strip()
        if not stripped:
            return True
        if stripped.endswith("{") or stripped.endswith("`") or stripped.endswith("if") or stripped.endswith("try"):
            return True
        return False

    @staticmethod
    def _python_has_semicolons(script: str) -> bool:
        return any(";" in line for line in script.splitlines())

    @staticmethod
    def _python_starts_with_import(script: str) -> bool:
        for line in script.splitlines():
            stripped = line.strip()
            if not stripped or stripped.startswith("#"):
                continue
            return stripped.startswith(("import ", "from "))
        return False

    @staticmethod
    def _python_uses_text_open_for_binary(script: str) -> bool:
        pattern = re.compile(r"open\([^\n]*\.(docx|xlsx|pdf)[^\n]*['\"]w", re.IGNORECASE)
        return bool(pattern.search(script))

    def _log_exec_result(self, result: dict[str, Any]) -> None:
        if not DEBUG:
            return
        stdout = (result.get("stdout") or "")[:1000]
        stderr = (result.get("stderr") or "")[:400]
        self._log_debug("SCRIPT_PATH", str(result.get("script_path")))
        self._log_debug("EXEC", str(result.get("exec_cmd")))
        if stdout:
            self._log_debug("STDOUT", stdout)
        if stderr:
            self._log_debug("STDERR_HEAD", stderr)
        self._log_debug("RC", f"{result.get('returncode')}")
        self._log_debug("DURATION_MS", f"{result.get('duration_ms')}")

    def run(self, user_input: str, stateless: bool = False) -> str:
        state = load_state()
        active_url = state.get("active_url")
        youtube_response = self._handle_youtube_command(user_input, active_url)
        if youtube_response:
            return youtube_response
        active_file = state.get("active_file")
        active_type = state.get("active_type")
        is_file_op = self._is_file_operation(user_input)
        has_explicit_file = self._has_explicit_file(user_input)

        if is_file_op and not has_explicit_file:
            if not active_file:
                response = "Какой файл изменить? Скажи /files и выбери /use 1 или напиши путь."
                if not stateless:
                    self.history.extend(
                        [
                            {"role": "user", "content": user_input},
                            {"role": "assistant", "content": response},
                        ]
                    )
                    self.history = self.history[-self.history_limit :]
                return response
            user_input = f"{self._format_active_context(active_file, active_type)}{user_input}"

        intent, forced_language = self._route_intent(user_input)
        if is_file_op and not has_explicit_file:
            forced_language = "python"
            intent = "file"
        if intent == "unknown":
            response = "Уточни, что именно нужно сделать на компьютере?"
            if not stateless:
                self.history.extend(
                    [
                        {"role": "user", "content": user_input},
                        {"role": "assistant", "content": response},
                    ]
                )
                self.history = self.history[-self.history_limit :]
            return response

        system_override = None
        if intent == "os_action":
            system_override = (
                "Ответь только одним блоком ```powershell``` и ничего больше. "
                "Никакого Python."
            )
        elif intent == "file":
            system_override = (
                "Ответь только одним блоком ```python``` и ничего больше. "
                "Для .docx используй python-docx, для .xlsx openpyxl, для .pdf reportlab. "
                "Запрещены однострочники с ';' и open(...,'w') для docx/xlsx/pdf. "
                "Первой строкой должны идти нужные импорты."
            )

        messages = self._build_messages(user_input, stateless, system_override)
        self.logger.log_user_input(user_input, len(messages))
        self.logger.log("user_input", {"content": user_input})
        request_confirmed = False
        last_language: str | None = None
        last_risk: RiskLevel | None = None
        risk_rank = {
            RiskLevel.LOW: 0,
            RiskLevel.MEDIUM: 1,
            RiskLevel.HIGH: 2,
        }

        for attempt in range(MAX_RETRIES + 1):
            response = self.client.chat(messages, tools=[], tool_choice="none")
            raw_content = response.choices[0].message.content or ""
            self._log_debug("LLM_RAW", sanitize_assistant_text(raw_content)[:300])
            language, script = self._extract_script(raw_content)
            if not language or not script:
                if self._is_command_text(raw_content):
                    language = "powershell"
                    script = raw_content.strip()
                else:
                    if forced_language:
                        messages.append(
                            {
                                "role": "user",
                                "content": (
                                    f"Ответь только кодом в блоке ```{forced_language}```."
                                ),
                            }
                        )
                        continue
                    assistant_text = sanitize_assistant_text(raw_content)
                    if not stateless:
                        self.history.extend(
                            [
                                {"role": "user", "content": user_input},
                                {"role": "assistant", "content": assistant_text},
                            ]
                        )
                        self.history = self.history[-self.history_limit :]
                    return assistant_text
            if forced_language and language != forced_language:
                messages.append(
                    {
                        "role": "user",
                        "content": (
                            f"Нужно вернуть только {forced_language} код одним блоком."
                        ),
                    }
                )
                continue
            self._log_debug("SCRIPT_LANG", language)
            self._log_debug("SCRIPT_HASH", self._script_hash(script))

            if language == "python":
                try:
                    compile(script, "<agent>", "exec")
                except SyntaxError as exc:
                    messages.append(
                        {
                            "role": "user",
                            "content": (
                                f"Скрипт синтаксически неверный: {exc}. "
                                "Верни исправленный полный python-код одним блоком."
                            ),
                        }
                    )
                    continue
                if self._python_has_semicolons(script):
                    messages.append(
                        {
                            "role": "user",
                            "content": (
                                "Нельзя использовать ';' в python. "
                                "Верни полный python-код без однострочников с ';'."
                            ),
                        }
                    )
                    continue
                if not self._python_starts_with_import(script):
                    messages.append(
                        {
                            "role": "user",
                            "content": (
                                "Python-скрипт должен начинаться с нужных импортов "
                                "(первая непустая строка — import/from)."
                            ),
                        }
                    )
                    continue
                if self._python_uses_text_open_for_binary(script):
                    messages.append(
                        {
                            "role": "user",
                            "content": (
                                "Нельзя создавать docx/xlsx/pdf через open(...,'w'). "
                                "Используй правильные библиотеки."
                            ),
                        }
                    )
                    continue
            if language == "powershell" and self._powershell_script_incomplete(script):
                messages.append(
                    {
                        "role": "user",
                        "content": (
                            "Скрипт PowerShell выглядит неполным или пустым. "
                            "Верни исправленный полный powershell-код одним блоком."
                        ),
                    }
                )
                continue

            current_risk = assess_risk(language, script)
            need_confirm = True
            if request_confirmed and last_language == language and last_risk:
                if risk_rank[current_risk] <= risk_rank[last_risk]:
                    need_confirm = False
            self._log_debug("CONFIRM", "repeat" if need_confirm else "reuse")

            if need_confirm:
                approved = confirm_if_needed(language, script)
                self.logger.log_policy("SCRIPT", "script_execution", approved)
                if not approved:
                    return "Ок, отменено пользователем."
                request_confirmed = True
                last_language = language
                last_risk = current_risk

            result = self._run_script(language, script)
            self._log_exec_result(result)
            stdout = (result.get("stdout") or "")[:2000]
            stderr = (result.get("stderr") or "")[:2000]
            ok = bool(result.get("ok"))
            returncode = result.get("returncode")

            if language == "python" and not ok:
                missing_module = self._extract_missing_module(stderr)
                if missing_module:
                    package = self._map_package(missing_module)
                    if package:
                        self._log_debug("MISSING_DEP", f"{missing_module} -> {package}")
                        install = input(f"Не установлен пакет {package}. Установить? (y/n): ").strip().lower()
                        if install == "y":
                            install_result = run_pip_install(package, TIMEOUT_SEC)
                            self._log_exec_result(install_result)
                            if install_result.get("ok"):
                                messages.append({"role": "user", "content": user_input})
                                continue

            if ok:
                summary = "✅ Готово"
                if stdout or stderr:
                    final_response = (
                        f"{summary}\nreturncode={returncode}\nstdout=\n{stdout}\nstderr=\n{stderr}"
                    )
                else:
                    final_response = f"{summary}\nreturncode={returncode}"
                if not stateless:
                    self.history.extend(
                        [
                            {"role": "user", "content": user_input},
                            {"role": "assistant", "content": final_response},
                        ]
                    )
                    self.history = self.history[-self.history_limit :]
                return final_response

            if attempt < MAX_RETRIES:
                error_message = (
                    "Скрипт упал. Вот результат выполнения:\n"
                    f"returncode={returncode}\n"
                    f"stdout=\n{stdout}\n"
                    f"stderr=\n{stderr}\n"
                    "Исправь скрипт и верни ТОЛЬКО новый код-блок."
                )
                messages.append({"role": "user", "content": error_message})
                continue

            final_error = (
                "❌ Ошибка выполнения\n"
                f"returncode={returncode}\nstdout=\n{stdout}\nstderr=\n{stderr}"
            )
            if not stateless:
                self.history.extend(
                    [
                        {"role": "user", "content": user_input},
                        {"role": "assistant", "content": final_error},
                    ]
                )
                self.history = self.history[-self.history_limit :]
            return final_error

        return ""
